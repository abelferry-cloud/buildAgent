"""Word document generation skill using python-docx."""

import os
from typing import Optional

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    Document = None

from agent.tools.base import Tool, ToolResult


class DocxGenerateTool(Tool):
    """Generate a Word document."""

    name = "docx_generate"
    description = "Generate a Word document with paragraphs and headings."

    def __init__(self, output_dir: str = "output"):
        super().__init__()
        self.output_dir = os.path.join(os.getcwd(), output_dir)
        os.makedirs(self.output_dir, exist_ok=True)

    def execute(
        self,
        title: str,
        paragraphs: list[str],
        output_filename: str = "document.docx",
    ) -> ToolResult:
        """Generate a DOCX file."""
        if Document is None:
            return ToolResult(
                output="",
                error="python-docx not installed. Run: pip install python-docx"
            )

        try:
            doc = Document()
            doc.add_heading(title, 0)

            for para in paragraphs:
                doc.add_paragraph(para)

            output_path = os.path.join(self.output_dir, output_filename)
            doc.save(output_path)
            return ToolResult(output=f"Word 文档已生成: {output_path}")
        except Exception as e:
            return ToolResult(output="", error=str(e))
